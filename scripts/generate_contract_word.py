"""
Word template ашиглаад хөдөлмөрийн гэрээ үүсгэх script
"""
import sys
import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("python-docx сан суугаагүй байна. Суулгаж байна...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document


def format_date_mongolian(date_str):
    """Огноог монгол хэлээр форматлах"""
    try:
        if isinstance(date_str, str):
            date = datetime.strptime(date_str, "%Y-%m-%d")
        else:
            date = date_str
        
        year = date.year
        month = date.month
        day = date.day
        
        return f"{year} оны {month} дугаар сарын {day}-ны өдөр"
    except Exception as e:
        print(f"Date formatting error: {e}")
        return str(date_str)


def generate_contract_word(contract_data, output_path):
    """
    Template файлыг ашиглаад Word гэрээ үүсгэх
    """
    script_dir = Path(__file__).parent.parent
    template_path = script_dir / "public" / "templates" / "contracts" / "template.docx"
    
    if not template_path.exists():
        raise FileNotFoundError(f"Template файл олдсонгүй: {template_path}")
    
    doc = Document(str(template_path))
    replacements = [
        ("2025 оны. . . . дугаар сарын ….-ны өдөр", format_date_mongolian(contract_data.get('startDate', datetime.now().strftime("%Y-%m-%d")))),
        ("№ .........", f"№ {contract_data.get('contractNumber', '')}"),
        ("Эрдэнэс-Тавантолгой ХК", contract_data.get('companyName', 'Эрдэнэс-Тавантолгой ХК')),
        (". . . . . . . . . . . . . . . овогтой. . . ............", f"{contract_data.get('employeeLastName', '')} овогтой {contract_data.get('employeeName', '')}"),
        (". . . . . . . . . . . . . . овогтой. . ............", f"{contract_data.get('employeeLastName', '')} овогтой {contract_data.get('employeeName', '')}"),
        ("Регистрийн дугаар: ................./", f"Регистрийн дугаар: {contract_data.get('registrationNumber', contract_data.get('employeeId', ''))}"),
        ("Регистрийн дугаар: .................", f"Регистрийн дугаар: {contract_data.get('registrationNumber', contract_data.get('employeeId', ''))}"),
        ("Албан тушаал: ...............", f"Албан тушаал: {contract_data.get('position', '')}"),
        ("Харьяалагдах нэгж: ..............", f"Харьяалагдах нэгж: {contract_data.get('department', '')}"),
        ("Үндсэн цалин: ................ /............................../-н төгрөг", f"Үндсэн цалин: {contract_data.get('salary', 0):,.0f} /{contract_data.get('salaryText', str(contract_data.get('salary', 0)))}/-н төгрөг"),
    ]
    
    def replace_in_text(text):
        """Replace function - зөвхөн нэг удаа replace хийх"""
        result = text
        
        for old_text, new_text in replacements:
            if old_text in result:
                result = result.replace(old_text, new_text, 1)
        
        if 'Гүйцэтгэх захирал' in result and re.search(r'Гүйцэтгэх захирал\s+\.{3,}', result):
            director_name = contract_data.get('directorName', '')
            if director_name and director_name.strip():
                result = re.sub(r'Гүйцэтгэх захирал\s+\.{3,}', 
                               f"Гүйцэтгэх захирал {director_name}", 
                               result, count=1)
            else:
                result = re.sub(r'Гүйцэтгэх захирал\s+\.{3,}', 
                               "Гүйцэтгэх захирал", 
                               result, count=1)
        if 'Ажлын цаг:' in result and re.search(r'Ажлын цаг:\s*\.{3,}', result):
            result = re.sub(r'Ажлын цаг:\s*\.{3,}', 
                           f"Ажлын цаг: {contract_data.get('workSchedule', 'Бүтэн цагийн (08:00-17:00)')}", 
                           result, count=1)
        if 'Гэрээний хугацаа:' in result and re.search(r'Гэрээний хугацаа:\s*\.{3,}', result):
            result = re.sub(r'Гэрээний хугацаа:\s*\.{3,}', 
                           f"Гэрээний хугацаа: {contract_data.get('contractDuration', 'Тодорхой хугацаагүй')}", 
                           result, count=1)
        if result.strip().startswith('. . .') and len(result.strip()) < 50 and not any(char.isalpha() for char in result.strip() if char not in '. '):
            employee_full_name = f"{contract_data.get('employeeLastName', '')} {contract_data.get('employeeName', '')}".strip()
            if employee_full_name:
                result = employee_full_name
        
        return result

    for para in doc.paragraphs:
        original_text = para.text
        new_text = replace_in_text(original_text)
        if original_text != new_text:
            para.text = new_text
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = replace_in_text(original_text)
                    if original_text != new_text:
                        para.text = new_text
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if len(para_text) > 0 and len(para_text) < 50:
            if para_text.startswith('. . .') and not any(char.isalpha() or char.isdigit() for char in para_text if char not in '. '):
                employee_full_name = f"{contract_data.get('employeeLastName', '')} {contract_data.get('employeeName', '')}".strip()
                if employee_full_name:
                    para.text = employee_full_name
    
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


if __name__ == "__main__":
    test_data = {
        'contractNumber': 'CT-2025-001',
        'employeeName': 'Баяр',
        'employeeLastName': 'Болд',
        'employeeId': 'EMP-2025-001',
        'registrationNumber': 'РД12345678',
        'position': 'Программист',
        'department': 'IT хэлтэс',
        'salary': 1500000,
        'salaryText': 'нэг сая таван зуун мянган',
        'startDate': '2025-01-15',
        'endDate': None,
        'contractType': 'FULL_TIME',
        'workSchedule': 'Бүтэн цагийн (08:00-17:00)',
        'contractDuration': '1 жил',
        'companyName': 'Эрдэнэс-Тавантолгой ХК',
        'directorName': 'Ж.Батбаярagamers',
        'city': 'Улаанбаатар хот'
    }
    
    output = generate_contract_word(test_data, "test_contract.docx")
    print(f"Гэрээ үүсгэгдлээ: {output}")
